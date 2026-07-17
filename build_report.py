"""
build_report.py
----------------
Generates the finished microproject report (AES_Microproject_Report.docx) from
the project results, plus a viva-prep sheet (AES_Viva_Prep.docx).

Sections 1-2 follow the submitted Stage-1 text (with the GloVe reference fixed);
Sections 3-4 and the references are written from the actual results in
outputs/ and figures/.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(ROOT, "figures")

HEADING_BLUE = RGBColor(0x1F, 0x38, 0x64)   # dark blue, like the template headings
ACCENT_BLUE = RGBColor(0x2E, 0x5B, 0x9A)


# --------------------------------------------------------------------------- #
# Small helpers
# --------------------------------------------------------------------------- #
def set_base_style(doc):
    """Times New Roman 11 as the document default."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def add_running_header(doc, text):
    """Put the programme name in the page header, like the template."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = text
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def heading(doc, text, size=15, color=HEADING_BLUE, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Times New Roman"
    return p


def body(doc, text, justify=True, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(11)
    return p


def add_figure(doc, path, caption, width=6.2):
    if not os.path.exists(path):
        body(doc, f"[missing figure: {path}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(h)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# =========================================================================== #
# MAIN REPORT
# =========================================================================== #
def build_report():
    doc = Document()
    set_base_style(doc)
    add_running_header(doc, "UG Program in Artificial Intelligence and Data Science")

    # ---------------- TITLE PAGE ----------------
    def center(text, size=12, bold=False, color=None, space=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    doc.add_paragraph("\n")
    center("Mahavir Education Trust's", 11)
    center("Shah & Anchor Kutchhi Engineering College", 16, bold=True,
           color=HEADING_BLUE)
    center("Autonomous Institute Affiliated to University of Mumbai", 9)
    doc.add_paragraph("\n")
    center("MICROPROJECT REPORT ON", 14, bold=True, color=ACCENT_BLUE)
    doc.add_paragraph("\n")
    center("Artificial Intelligence", 26, bold=True, color=HEADING_BLUE)
    doc.add_paragraph("\n\n")
    center("Course Name: Artificial Intelligence (ADCR1503)", 12)
    center("Microproject Title: Automatic Essay Scoring", 12)
    doc.add_paragraph("\n")
    center("Submitted by (Group of 2):", 12, bold=True)
    center("1. Name: ______________________    Roll No: __________", 12)
    center("2. Name: ______________________    Roll No: __________", 12)
    center("Class: ______________________", 12)
    doc.add_paragraph("\n\n")
    center("Submitted to:", 12, bold=True)
    center("Department of Artificial Intelligence and Data Science", 12, bold=True)
    center("Shah and Anchor Kutchhi Engineering College, Chembur - 400088", 11)
    center("Academic Year: 2025-26", 11, bold=True)
    doc.add_page_break()

    # ---------------- ABSTRACT ----------------
    heading(doc, "ABSTRACT")
    body(doc,
         "Manual grading of essays is time-consuming, expensive, and prone to "
         "inconsistency between different human raters. This microproject applies "
         "deep learning to the problem of Automatic Essay Scoring (AES), where a "
         "model learns to predict a numeric score for a student essay based on its "
         "content, coherence, and language quality. The study covers the AES domain "
         "problem, dataset collection and preprocessing, the selection and design of "
         "a suitable deep learning architecture, and a full training and evaluation "
         "pipeline benchmarked against a handcrafted-feature baseline. A Bidirectional "
         "LSTM with mean-over-time pooling and GloVe word embeddings is trained on the "
         "ASAP dataset. A single BiLSTM trained across all eight essay prompts achieves "
         "a mean per-set Quadratic Weighted Kappa (QWK) of 0.686, matching an ensemble "
         "of eight prompt-specific linear baselines (0.672) and greatly exceeding a "
         "single global baseline (0.361). The project also considers deployment "
         "challenges, ethical concerns such as scoring bias, and real-world "
         "applications in education and standardized testing.")

    # ---------------- 1. INTRODUCTION ----------------
    heading(doc, "1. INTRODUCTION")
    body(doc,
         "Automatic Essay Scoring (AES) is the task of using a computational model to "
         "assign a score to a student essay in place of, or alongside, a human rater. "
         "Traditional grading is slow, costly to scale, and often inconsistent, since "
         "two human graders can score the same essay differently. Deep Learning is "
         "well suited to this problem because it can learn rich representations of "
         "language directly from text, capturing aspects such as vocabulary usage, "
         "sentence structure, coherence, and argument organization, without relying "
         "entirely on hand-crafted features. The motivation behind this project is to "
         "design a deep learning pipeline that reads an essay and predicts a score "
         "that closely matches the score a human rater would give, and to understand "
         "the practical constraints involved: essays vary widely in length, topic, and "
         "writing style; training data is limited and expensive to obtain; and scores "
         "must remain fair and explainable rather than being an opaque black box.")

    body(doc, "Objectives:", bold=True, justify=False)
    for t in [
        "To analyze the Automatic Essay Scoring domain problem and identify its key "
        "constraints and objectives.",
        "To collect and preprocess a suitable essay-scoring dataset for model training.",
        "To select and design a deep learning architecture capable of learning "
        "scoring patterns from essay text.",
        "To plan a training, validation, and evaluation strategy, and compare the "
        "proposed model against a baseline approach.",
    ]:
        bullet(doc, t)

    body(doc, "Scope:", bold=True, justify=False)
    body(doc,
         "This project focuses on the Natural Language Processing (NLP) domain, "
         "specifically text-based regression/scoring. The insights gained are "
         "applicable to ed-tech platforms, standardized testing (e.g., TOEFL, "
         "GRE-style essay sections), classroom assessment tools, and any system that "
         "needs to rate the quality of written text at scale.")

    # ---------------- 2. METHODOLOGY ----------------
    heading(doc, "2. METHODOLOGY")
    body(doc,
         "The methodology adopted in this project includes data collection, "
         "preprocessing, model selection, training, and evaluation. The workflow "
         "involves the following steps:")

    steps = [
        ("Dataset Description:",
         " The Automated Student Assessment Prize (ASAP) essay dataset (available on "
         "Kaggle) is used. It contains 12,976 student essays across 8 essay sets, each "
         "hand-scored by human raters, with scores on different scales per set "
         "(e.g., Set 3 uses 0-3 while Set 8 uses 0-60)."),
        ("Data Preprocessing:",
         " Text cleaning (lowercasing, collapsing the ASAP anonymisation tags such as "
         "@PERSON1 to placeholder words, removing special characters), tokenization, "
         "padding/truncation to a fixed sequence length, converting words to 50-"
         "dimensional GloVe embeddings, and normalizing each essay's domain1_score to "
         "a common 0-1 range using the OFFICIAL per-set ASAP score range (not the "
         "observed min/max). Data is split into training, validation, and test sets, "
         "stratified by essay set and score bin."),
        ("Model Architecture:",
         " A Bidirectional LSTM (BiLSTM) network with a GloVe-initialised embedding "
         "layer, a BiLSTM layer that returns an output at every timestep, mean-over-"
         "time pooling that averages those outputs across the essay, dropout, and a "
         "dense regression head with a sigmoid output that predicts the normalised "
         "score."),
        ("Training Process:",
         " Adam optimizer, Mean Squared Error (MSE) loss (scoring is framed as "
         "regression), batch size of 32, and up to 30 epochs with early stopping on "
         "validation loss and restoration of the best-epoch weights."),
        ("Evaluation:",
         " RMSE, MAE, and Quadratic Weighted Kappa (QWK) - the standard AES metric "
         "that measures agreement with human raters. QWK is computed on scores "
         "rescaled back to the original integer range. The model is compared against a "
         "baseline linear regression built on handcrafted features (word count, "
         "sentence count, average word length, unique-word ratio, spelling-error "
         "count), and QWK is reported PER SET as well as overall."),
        ("Deployment and Ethical Considerations:",
         " Key challenges include ensuring the model generalizes to unseen essay "
         "topics, avoiding bias toward superficial features (e.g., essay length) "
         "rather than genuine writing quality, keeping predictions explainable, and "
         "providing a human-in-the-loop review process for borderline or low-"
         "confidence scores."),
    ]
    for i, (title, text) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        p.add_run(f"{i}. ").bold = True
        p.add_run(title).bold = True
        p.add_run(text)

    # 2.1 Individual contributions
    heading(doc, "2.1 Individual Contribution to Design Methodology", size=13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Student 1 (Name: ____________, Roll No: ______): ").bold = True
    p.add_run(
        "Proposed and prepared the ASAP dataset and implemented the data pipeline "
        "(Stage A data exploration and Stage B preprocessing): loading the TSV with "
        "the correct ISO-8859-1 encoding, text cleaning including collapsing the ASAP "
        "anonymisation tags, tokenisation fitted on the training split only, and "
        "padding/truncation. Made and justified the key decision to normalise scores "
        "per essay set using the official ASAP ranges rather than the observed "
        "min/max. Designed and implemented the BiLSTM model (Stage D) and ran the "
        "architecture-tuning study — introducing mean-over-time pooling and moving "
        "from a frozen to a trainable GloVe embedding — which raised the set-1 QWK "
        "from 0.30 to 0.70.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Student 2 (Name: ____________, Roll No: ______): ").bold = True
    p.add_run(
        "Defined the evaluation methodology and implemented the metrics module — RMSE, "
        "MAE, and Quadratic Weighted Kappa computed on denormalised integer scores. "
        "Built the handcrafted-feature linear-regression baseline (Stage C) and "
        "designed the fair comparison: single global baseline versus eight per-prompt "
        "baselines, and per-set versus pooled QWK. Identified that pooled QWK across "
        "differently-scaled prompts is misleading and introduced score-bin "
        "stratification for representative splits. Produced the Stage E result figures "
        "and researched the deployment and ethical considerations for AES.")

    # 2.2 Group consensus
    heading(doc, "2.2 Design Methodology Quality (Group Consensus)", size=13)
    body(doc,
         "After discussing individual proposals, the group agreed on the BiLSTM-based "
         "architecture with the ASAP dataset, MSE loss with QWK/RMSE/MAE as evaluation "
         "metrics, and a handcrafted-feature linear-regression baseline for comparison. "
         "This methodology was chosen because it balances implementation feasibility "
         "within the project timeline with strong alignment to established AES research "
         "practice, and both members' proposals were integrated into the final design "
         "(dataset, preprocessing and architecture from Student 1; evaluation, baseline "
         "and ethics from Student 2).")

    doc.add_page_break()

    # ---------------- 3. IMPLEMENTATION AND RESULTS ----------------
    heading(doc, "3. IMPLEMENTATION AND RESULTS")
    body(doc,
         "The project was implemented in Python using TensorFlow/Keras, scikit-learn, "
         "pandas and NLTK, following the methodology in Section 2. The code is "
         "organised into five stages (A-E): data exploration, preprocessing, baseline "
         "model, BiLSTM model, and evaluation. A fixed random seed (42) is set across "
         "Python, NumPy and TensorFlow so every result is reproducible.")

    body(doc, "3.1 Dataset Exploration", bold=True, justify=False)
    body(doc,
         "The ASAP dataset contains 12,976 essays across 8 prompts, with no missing "
         "target scores. The number of essays is fairly balanced across sets (Set 8 is "
         "smallest at 723). Crucially, the score scales differ dramatically between "
         "prompts, which is why per-set normalisation is required before a single "
         "model can learn from all prompts together.")
    add_figure(doc, os.path.join(FIG, "A_essays_per_set.png"),
               "Figure 1: Number of essays per essay set.")
    add_figure(doc, os.path.join(FIG, "A_score_ranges_per_set.png"),
               "Figure 2: Score distribution per set — note the very different scales "
               "(1-6 up to 0-60), motivating per-set normalisation.")
    add_figure(doc, os.path.join(FIG, "A_length_distribution.png"),
               "Figure 3: Essay-length distribution across all sets, used to choose "
               "the padding/truncation length.")

    body(doc, "3.2 Models and Training", bold=True, justify=False)
    body(doc,
         "The baseline is a linear regression on five handcrafted features (word "
         "count, sentence count, average word length, unique-word ratio, spelling-"
         "error count); its learned weights are dominated by word count, i.e. it "
         "essentially rewards longer essays. The deep model is a BiLSTM: a GloVe-50d "
         "embedding (77.4% of the vocabulary matched a pretrained vector), a "
         "bidirectional LSTM, mean-over-time pooling, dropout, and a dense sigmoid "
         "output. It was trained with Adam and MSE loss, batch size 32, and early "
         "stopping with best-weight restoration.")
    body(doc,
         "An architecture-tuning study on essay set 1 showed how each design choice "
         "contributed, which is itself an instructive result:")
    add_table(doc,
              ["Design configuration (essay set 1)", "QWK"],
              [["Frozen embedding + final LSTM state only", "0.30"],
               ["+ mean-over-time pooling", "0.47"],
               ["+ trainable embedding, reduced dropout", "0.70"]])

    body(doc, "3.3 Results", bold=True, justify=False)
    body(doc,
         "The headline metric is the MEAN PER-SET QWK. A single BiLSTM trained on all "
         "eight prompts matches an ensemble of eight separately-tuned per-prompt "
         "baselines, and is far better than a single global feature-based model:")
    add_table(doc,
              ["Model", "Mean per-set QWK"],
              [["BiLSTM — one model, all 8 sets", "0.686"],
               ["Baseline — 8 separate per-prompt models", "0.672"],
               ["Baseline — one global model, all 8 sets", "0.361"]])

    body(doc,
         "The per-set breakdown (Table 4, Figure 4) shows a genuinely even split: the "
         "BiLSTM wins on four prompts and the strongest baseline wins on four. The "
         "baseline edges the longest, most length-driven prompt (Set 1), while the "
         "BiLSTM wins where content matters more (Sets 4, 6, 7).")
    add_table(doc,
              ["Set", "Baseline QWK", "BiLSTM QWK", "BiLSTM RMSE", "BiLSTM MAE", "Winner"],
              [["1", "0.790", "0.723", "0.970", "0.762", "Baseline"],
               ["2", "0.580", "0.611", "0.566", "0.464", "BiLSTM"],
               ["3", "0.677", "0.633", "0.583", "0.448", "Baseline"],
               ["4", "0.689", "0.771", "0.540", "0.450", "BiLSTM"],
               ["5", "0.814", "0.766", "0.525", "0.418", "Baseline"],
               ["6", "0.620", "0.774", "0.526", "0.435", "BiLSTM"],
               ["7", "0.654", "0.719", "2.968", "2.324", "BiLSTM"],
               ["8", "0.554", "0.493", "4.904", "3.786", "Baseline"],
               ["Mean", "0.672", "0.686", "—", "—", "—"]])

    add_figure(doc, os.path.join(FIG, "E1_qwk_per_set_bars.png"),
               "Figure 4: Per-set QWK, baseline vs BiLSTM, with each model's mean.")
    add_figure(doc, os.path.join(FIG, "E2_loss_curves.png"),
               "Figure 5: Training vs validation loss. The set-1 run overfits (wide "
               "train/val gap); the all-sets run converges more healthily with 7x more "
               "data.")

    body(doc,
         "A note on metrics: pooling all eight prompts into a single QWK gives an "
         "inflated value of 0.98, which is misleading. Because the prompts use very "
         "different score ranges, essay pairs from different sets trivially 'agree' on "
         "being far apart, so the pooled metric partly rewards the model for sorting "
         "which set an essay belongs to rather than scoring within a set. Figure 6 "
         "makes this visible as eight separated clusters. Mean per-set QWK is therefore "
         "the correct metric and is used throughout.")
    add_figure(doc, os.path.join(FIG, "E3_pred_vs_actual_pooled.png"),
               "Figure 6: Predicted vs actual, pooled and coloured by set. The eight "
               "separated clusters are exactly why a single pooled QWK is inflated.")
    add_figure(doc, os.path.join(FIG, "E4_pred_vs_actual_per_set.png"),
               "Figure 7: BiLSTM predicted vs actual, per set, with QWK annotated.")

    # ---------------- 4. CONCLUSION ----------------
    heading(doc, "4. CONCLUSION AND FUTURE SCOPE")
    body(doc,
         "This microproject set out to determine not only whether a deep learning model "
         "can score essays, but whether it is genuinely suitable for the task compared "
         "with a simple baseline. The answer is a qualified yes. A single BiLSTM that "
         "reads only the essay text achieves a mean per-set QWK of 0.686 — matching an "
         "ensemble of eight separately-tuned per-prompt baselines (0.672) and far "
         "exceeding a single global feature model (0.361). The practical significance "
         "is a deployment argument: one neural model can serve all eight prompts to the "
         "same standard that the baseline reaches only by maintaining eight separate "
         "specialised models. In a real ed-tech system, one model to maintain, monitor "
         "and re-train is substantially cheaper and safer than eight.")
    body(doc,
         "The per-set comparison is deliberately honest rather than a single headline "
         "number. The four-to-four split between the models shows that surface features "
         "such as length remain a strong signal for some prompt types — notably Set 1, "
         "the longest and most structured prompt, where the baseline wins — while the "
         "BiLSTM's ability to read content pays off on prompts where quality is less "
         "correlated with length. The weakest result, Set 8 (QWK 0.49), is named as a "
         "genuine limitation: it is the smallest training set with the widest score "
         "scale (0-60) and the fewest test essays, making it both the hardest to learn "
         "and the noisiest to measure.")
    body(doc,
         "Future work. The model can be improved with larger and transfer-learning "
         "architectures such as a fine-tuned BERT regressor; with per-set output heads "
         "or ordinal-regression losses better suited to wide score scales; with a "
         "hierarchical encoder to avoid truncating very long essays (as in Set 8); and "
         "with data augmentation or additional graded essays for the smaller prompts. "
         "More robust evaluation via cross-validation would also reduce the measurement "
         "noise seen on the smaller sets.")

    # ---------------- REFERENCES ----------------
    heading(doc, "REFERENCES")
    refs = [
        "Ian Goodfellow, Yoshua Bengio, and Aaron Courville, Deep Learning, MIT Press, "
        "2016.",
        "Francois Chollet, Deep Learning with Python, Manning Publications, 2018.",
        "Aurelien Geron, Hands-On Machine Learning with Scikit-Learn, Keras & "
        "TensorFlow, O'Reilly Media, 2022.",
        "TensorFlow Documentation: https://www.tensorflow.org/",
        "scikit-learn Documentation: https://scikit-learn.org/",
        "Hewlett Foundation, Automated Student Assessment Prize (ASAP) - Essay Scoring "
        "Dataset (including the official per-set score ranges), Kaggle, 2012. "
        "https://www.kaggle.com/c/asap-aes",
        "Jeffrey Pennington, Richard Socher, and Christopher D. Manning, GloVe: Global "
        "Vectors for Word Representation. 2024 Wikipedia + Gigaword 50-dimensional "
        "release. Stanford NLP. https://nlp.stanford.edu/projects/glove/",
        "Taghipour, K. and Ng, H.T., A Neural Approach to Automated Essay Scoring, "
        "Proceedings of EMNLP, 2016.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{i}. ").bold = True
        p.add_run(r).font.size = Pt(10.5)

    out = os.path.join(ROOT, "AES_Microproject_Report.docx")
    doc.save(out)
    return out


# =========================================================================== #
# VIVA PREP SHEET
# =========================================================================== #
def build_viva():
    doc = Document()
    set_base_style(doc)
    heading(doc, "Automatic Essay Scoring — Viva Preparation", size=16)
    body(doc, "Likely questions and prepared answers.", italic=True)

    qa = [
        ("1. Why does mean per-set QWK matter more than pooled QWK?",
         "QWK measures chance-corrected agreement, and it penalises disagreements more "
         "the further apart they are. When you pool predictions from eight prompts that "
         "use very different score ranges (Set 3 is 0-3, Set 8 is 0-60), an essay scored "
         "3 and one scored 55 can never be 'close', so cross-set pairs trivially agree "
         "on being far apart. The pooled metric then partly rewards the model for "
         "sorting which SET an essay belongs to, not how well it scores WITHIN a set — "
         "which inflates it to 0.98. Mean per-set QWK only ever compares essays on the "
         "same scale, so 0.686 is the honest number."),
        ("2. Why did the baseline beat the BiLSTM on the set-1-only run but lose on the "
         "combined run?",
         "On set 1 alone there are only ~1,250 training essays. The length-dominated "
         "linear baseline is very strong there (QWK 0.79), while the BiLSTM overfits on "
         "so little data — its training and validation loss diverge — reaching only "
         "0.72. On the combined run the BiLSTM sees ~9,000 essays and, as one model, "
         "learns to generalise across prompts. A single global baseline cannot fit the "
         "eight different feature-to-score relationships with one set of weights, so it "
         "collapses to 0.361. The fair comparison — eight separate per-prompt baselines "
         "— ties the BiLSTM (0.672 vs 0.686). So it comes down to data scale and whether "
         "the baseline is allowed to specialise per prompt."),
        ("3. Why use the official score ranges instead of the observed min/max for "
         "normalisation?",
         "Min-max scaling on observed data assumes no future essay can score below or "
         "above what happened to appear in this training file. If a test essay is a "
         "genuine 0 or a genuine 30 on Set 7, observed-range scaling would map it "
         "outside [0,1] and corrupt both training and the QWK rescaling. The official "
         "ASAP ranges are documented, fixed, and defensible, and they keep the shared "
         "0-1 target meaningful and comparable across all prompts. It is also standard "
         "practice in published AES work."),
        ("4. What would you try next to close the gap on Set 8?",
         "Set 8 is the smallest set (505 training essays) with the widest scale (0-60) "
         "and only 109 test essays, so it is both hard to learn and noisy to measure. I "
         "would (a) get more graded essays or use data augmentation for it; (b) avoid "
         "truncation of its long (600+ word) essays with a longer window or a "
         "hierarchical sentence-then-document encoder; (c) use a per-set output head or "
         "an ordinal-regression loss suited to the wide scale; (d) try transfer learning "
         "with a fine-tuned BERT; and (e) evaluate with cross-validation to reduce the "
         "measurement noise from the tiny test set."),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(q)
        r.bold = True
        r.font.color.rgb = ACCENT_BLUE
        body(doc, a)

    out = os.path.join(ROOT, "AES_Viva_Prep.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    r1 = build_report()
    r2 = build_viva()
    print("Saved:", r1)
    print("Saved:", r2)
